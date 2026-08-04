"""Tests for PDF resume parsing."""

from pathlib import Path
import tempfile
import unittest

import fitz
from docx import Document

from src.parser import read_resumes


class ResumeParserTests(unittest.TestCase):
    def _create_pdf(self, path: Path, text: str) -> None:
        document = fitz.open()
        page = document.new_page()
        page.insert_text((72, 72), text)
        document.save(path)
        document.close()

    def _create_docx(self, path: Path) -> None:
        document = Document()
        document.add_paragraph("Alice Doe")
        document.add_paragraph("Python Developer")
        table = document.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "alice@example.com"
        document.save(path)

    def test_read_resumes_returns_cleaned_text_for_each_pdf(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            resume_folder = Path(directory)
            self._create_pdf(resume_folder / "alice.pdf", "Alice\n\nPython Developer")
            self._create_pdf(resume_folder / "john.pdf", "John\tData Engineer")
            (resume_folder / "notes.txt").write_text("Ignore me", encoding="utf-8")

            resume_texts = read_resumes(resume_folder)

        self.assertEqual(
            resume_texts,
            {
                "alice.pdf": "Alice\nPython Developer",
                "john.pdf": "John Data Engineer",
            },
        )

    def test_read_resumes_rejects_missing_folder(self) -> None:
        with self.assertRaises(FileNotFoundError):
            read_resumes("missing-resume-folder")

    def test_read_resumes_reads_docx_paragraphs_and_tables(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            resume_folder = Path(directory)
            self._create_docx(resume_folder / "alice.DOCX")

            resume_texts = read_resumes(resume_folder)

        self.assertEqual(
            resume_texts,
            {"alice.DOCX": "Alice Doe\nPython Developer\nalice@example.com"},
        )


if __name__ == "__main__":
    unittest.main()
