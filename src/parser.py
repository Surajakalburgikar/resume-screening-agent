"""Read and clean resume text from PDF and DOCX files."""

from pathlib import Path
import re

import fitz
from docx import Document


SUPPORTED_RESUME_EXTENSIONS = {".pdf", ".docx"}


def clean_text(text: str) -> str:
    """Normalize inline whitespace while preserving meaningful line breaks."""
    cleaned_lines = [
        re.sub(r"[\t\f\v ]+", " ", line).strip()
        for line in text.splitlines()
    ]
    return "\n".join(line for line in cleaned_lines if line)


def _read_pdf(path: Path) -> str:
    with fitz.open(path) as document:
        return "".join(page.get_text() for page in document)


def _read_docx(path: Path) -> str:
    document = Document(path)
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    table_cells = [
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    ]
    return "\n".join([*paragraphs, *table_cells])


def read_resumes(resume_folder: str | Path) -> dict[str, str]:
    """Return cleaned text for each PDF or DOCX resume in ``resume_folder``.

    Files are processed in name order to keep downstream rankings reproducible.
    The returned dictionary uses each file name (for example, ``john.pdf``) as
    its key. Unsupported files are ignored.
    """
    folder = Path(resume_folder)
    if not folder.is_dir():
        raise FileNotFoundError(f"Resume folder does not exist: {folder}")

    resume_files = sorted(
        (
            path
            for path in folder.iterdir()
            if path.is_file() and path.suffix.lower() in SUPPORTED_RESUME_EXTENSIONS
        ),
        key=lambda path: path.name.lower(),
    )
    resume_texts: dict[str, str] = {}
    for resume_path in resume_files:
        extracted_text = (
            _read_pdf(resume_path)
            if resume_path.suffix.lower() == ".pdf"
            else _read_docx(resume_path)
        )
        resume_texts[resume_path.name] = clean_text(extracted_text)

    return resume_texts
